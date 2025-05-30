import io
import datetime
from django.views.decorators.http import require_POST
from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from lessons.models import Lesson, LessonProgress, GRADING_PERIOD_CHOICES as LESSON_GRADING_PERIOD_CHOICES
from quizzes.models import Quiz, QuizAttempt
from exams.models import Exam, ExamAttempt
from .models import StudentProgress, ActivityLog, Attendance
from django.db.models import Q, Sum, Avg, Count
from django.utils import timezone
from datetime import timedelta
from django.template.defaultfilters import floatformat
from docx import Document
import docx.shared
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from django.contrib.auth import get_user_model
import os
from django.conf import settings
from django.http import JsonResponse

TEMPLATE_PATH = "media/templates/RANKING_TEMPLATE.docx"

# Use imported choices if available, otherwise fallback to default
GRADING_PERIODS = LESSON_GRADING_PERIOD_CHOICES if 'LESSON_GRADING_PERIOD_CHOICES' in locals() else [
    ('prelim', 'Prelim'),
    ('midterm', 'Midterm'),
    ('prefinal', 'Prefinal'),
    ('final', 'Final'),
]

def get_period_label(period_value):
    for val, label in GRADING_PERIODS:
        if val == period_value:
            return label
    return period_value

def has_grading_period_field(model):
    return 'grading_period' in [f.name for f in model._meta.get_fields()]

@login_required
def dashboard(request):
    if request.user.is_staff:
        return admin_dashboard(request)
    return student_dashboard(request)

@login_required
def student_dashboard(request):
    user = request.user

    grading_periods = GRADING_PERIODS
    period_stats = {}

    quiz_points_qs = QuizAttempt.objects.filter(
        user=user,
        completed=True,
        quiz__is_archived=False 
    )
    exam_points_qs = ExamAttempt.objects.filter(user=user, completed=True)
    total_quiz_points = quiz_points_qs.aggregate(
        s=Sum('raw_points')
    )['s']
    if total_quiz_points is None:
        total_quiz_points = quiz_points_qs.aggregate(s=Sum('score'))['s'] or 0

    total_exam_points = exam_points_qs.aggregate(
        s=Sum('raw_points')
    )['s']
    if total_exam_points is None:
        total_exam_points = exam_points_qs.aggregate(s=Sum('score'))['s'] or 0

    from django.contrib.auth import get_user_model
    User = get_user_model()
    # Get filter parameters
    selected_period = request.GET.get('grading_period', 'overall')
    selected_section = request.GET.get('section', 'all')
    search_query = request.GET.get('search', '').strip()
    
    # Initial queryset: all non-staff students
    students_qs = User.objects.filter(is_staff=False)
    
    # Apply section filter if available
    if selected_section != 'all':
        # Attempt to join with students or from User if 'section' field
        try:
            from students.models import Student
            student_ids = Student.objects.filter(section=selected_section).values_list('user_id', flat=True)
            students_qs = students_qs.filter(id__in=student_ids)
        except Exception:
            # Fallback: check User.section if available
            if hasattr(User, 'section'):
                students_qs = students_qs.filter(section=selected_section)
    
    # Apply search filter
    if search_query:
        students_qs = students_qs.filter(
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(username__icontains=search_query)
        )
    
    # Define all_students as the filtered queryset
    all_students = students_qs
    
    student_totals = []
    for student in all_students:
        quiz_qs = QuizAttempt.objects.filter(
            user=student,
            completed=True,
            quiz__is_archived=False  
        )
        exam_qs = ExamAttempt.objects.filter(user=student, completed=True)
        stu_quiz_points = quiz_qs.aggregate(s=Sum('raw_points'))['s']
        if stu_quiz_points is None:
            stu_quiz_points = quiz_qs.aggregate(s=Sum('score'))['s'] or 0
        stu_exam_points = exam_qs.aggregate(s=Sum('raw_points'))['s']
        if stu_exam_points is None:
            stu_exam_points = exam_qs.aggregate(s=Sum('score'))['s'] or 0
        student_totals.append({
            'user_id': student.id,
            'total_points': (stu_quiz_points or 0) + (stu_exam_points or 0)
        })
    sorted_students = sorted(student_totals, key=lambda x: x['total_points'], reverse=True)
    overall_rank = None
    for idx, stu in enumerate(sorted_students, 1):
        if stu['user_id'] == user.id:
            overall_rank = idx
            break
    total_students = all_students.count()

    for period_value, period_label in grading_periods:
        lessons_in_period = Lesson.objects.filter(grading_period=period_value, is_active=True, is_archived=False)
        total_lessons = lessons_in_period.count()
        lessons_read_count = LessonProgress.objects.filter(
            user=user,
            lesson__grading_period=period_value,
            completed=True,
            lesson__is_active=True,
            lesson__is_archived=False
        ).count() if total_lessons > 0 else 0

        lessons_progress_percent = round(lessons_read_count / total_lessons * 100, 1) if total_lessons else 0

        quizzes_answered_count = 0
        quiz_result_list = []
        if has_grading_period_field(QuizAttempt):
            quiz_q = Q(user=user, completed=True, grading_period=period_value)
            quizzes_attempts_qs = QuizAttempt.objects.filter(quiz_q, quiz__is_archived=False).select_related('quiz')
            quizzes_answered_count = quizzes_attempts_qs.count()
            for qa in quizzes_attempts_qs:
                all_users_attempts = QuizAttempt.objects.filter(
                    quiz=qa.quiz,
                    completed=True,
                    grading_period=period_value,
                    quiz__is_archived=False
                ).order_by('-score')
                all_scores = list(all_users_attempts.values_list('score', flat=True))
                placer = all_scores.index(qa.score) + 1 if qa.score in all_scores else None
                top_score = all_scores[0] if all_scores else None
                badge = placer == 1 and qa.score == top_score
                earned_points = getattr(qa, "raw_points", None)
                total_points = getattr(qa, "total_points", None)
                max_score = getattr(qa.quiz, 'max_score', 100)
                score_percent = round((qa.score / max_score) * 100, 1) if max_score else 0
                quiz_result_list.append({
                    'quiz_title': qa.quiz.title,
                    'score': qa.score,
                    'total_score': max_score,
                    'earned_points': earned_points,
                    'total_points': total_points,
                    'score_percent': score_percent,
                    'placer': placer,
                    'badge': badge,
                    'raw_score': qa.score,  
                })

        exam_result_list = []
        if has_grading_period_field(ExamAttempt):
            exam_q = Q(user=user, completed=True, grading_period=period_value)
            exams_attempts_qs = ExamAttempt.objects.filter(exam_q).select_related('exam')

            for ea in exams_attempts_qs:
                all_users_attempts = ExamAttempt.objects.filter(
                    exam=ea.exam,
                    completed=True,
                    grading_period=period_value
                ).order_by('-score')
                all_scores = list(all_users_attempts.values_list('score', flat=True))
                placer = all_scores.index(ea.score) + 1 if ea.score in all_scores else None
                top_score = all_scores[0] if all_scores else None
                badge = placer == 1 and ea.score == top_score
                earned_points = getattr(ea, "raw_points", None)
                total_points = getattr(ea, "total_points", None)
                max_score = getattr(ea.exam, 'max_score', 100)
                score_percent = round((ea.score / max_score) * 100, 1) if max_score else 0
                status = "passed" if getattr(ea, "passed", False) else "failed"
                exam_result_list.append({
                    'exam_title': ea.exam.title,
                    'score': ea.score,
                    'total_score': max_score,
                    'earned_points': earned_points,
                    'total_points': total_points,
                    'score_percent': score_percent,
                    'placer': placer,
                    'status': status,
                    'badge': badge,
                    'raw_score': ea.score,  
                })

        period_stats[period_value] = {
            'label': period_label,
            'lessons_read_count': lessons_read_count,
            'total_lessons': total_lessons,
            'lessons_progress_percent': lessons_progress_percent,
            'quizzes_answered_count': quizzes_answered_count,
            'quiz_result_list': quiz_result_list,
            'exam_result_list': exam_result_list
        }


    lessons_achievements = []
    lessons_read_period = {}
    for period_value, period_label in grading_periods:
        active_lessons = Lesson.objects.filter(grading_period=period_value, is_active=True, is_archived=False).count()
        lessons_read = LessonProgress.objects.filter(
            user=user,
            lesson__grading_period=period_value,
            completed=True,
            lesson__is_active=True,
            lesson__is_archived=False
        ).count() if active_lessons > 0 else 0

        lessons_read_period[period_value] = (lessons_read, active_lessons)
        if active_lessons > 0 and lessons_read == active_lessons:
            lessons_achievements.append({
                'type': 'finished_lessons_period',
                'description': f'100% lessons read at {period_label}',
                'icon': 'book'
            })

    active_total_lessons = Lesson.objects.filter(is_active=True, is_archived=False).count()
    all_read = LessonProgress.objects.filter(
        user=user,
        lesson__is_active=True,
        lesson__is_archived=False,
        completed=True
    ).values('lesson_id').distinct().count()

    if active_total_lessons > 0 and all_read == active_total_lessons:
        lessons_achievements.append({
            'type': 'finished_lessons_all',
            'description': f'100% lessons read (All)',
            'icon': 'bookshelf'
        })

    quiz_achievements = []
    user_quiz_attempts = QuizAttempt.objects.filter(
        user=user,
        completed=True,
        quiz__is_archived=False
    ).select_related('quiz')
    for qa in user_quiz_attempts:
        has_perfect = False
        if hasattr(qa, "raw_points") and hasattr(qa, "total_points") and qa.raw_points is not None and qa.total_points is not None:
            if qa.raw_points == qa.total_points and qa.total_points > 0:
                has_perfect = True
        else:
            max_score = getattr(qa.quiz, 'max_score', 100)
            if qa.score == max_score and max_score > 0:
                has_perfect = True
            elif qa.score >= 100:
                has_perfect = True
        if has_perfect:
            quiz_achievements.append({
                'quiz_title': qa.quiz.title,
                'description': f'Perfect score at {qa.quiz.title}',
                'icon': 'check-circle'
            })

    exam_achievements = []
    user_exam_attempts = ExamAttempt.objects.filter(user=user, completed=True).select_related('exam')
    for ea in user_exam_attempts:
        has_perfect = False
        period = getattr(ea, 'grading_period', None)
        period_label = get_period_label(period)
        if hasattr(ea, "raw_points") and hasattr(ea, "total_points") and ea.raw_points is not None and ea.total_points is not None:
            if ea.raw_points == ea.total_points and ea.total_points > 0:
                has_perfect = True
        else:
            max_score = getattr(ea.exam, 'max_score', 100)
            if ea.score == max_score and max_score > 0:
                has_perfect = True
            elif ea.score >= 100:
                has_perfect = True
        if has_perfect:
            exam_achievements.append({
                'exam_title': ea.exam.title,
                'description': f'Perfect score at {ea.exam.title} ({period_label})',
                'icon': 'star-fill'
            })



    context = {
        'grading_periods': grading_periods,
        'period_stats': period_stats,
        'overall_rank': overall_rank,
        'total_students': total_students,
        'total_quiz_points': total_quiz_points,
        'total_exam_points': total_exam_points,
        'lessons_achievements': lessons_achievements,
        'quiz_achievements': quiz_achievements,
        'exam_achievements': exam_achievements,
    }
    return render(request, 'dashboard/student.html', context)

@login_required
def admin_dashboard(request):
    if not request.user.is_staff:
        messages.error(request, "You don't have permission to view this page.")
        return redirect('student_dashboard')

    from django.contrib.auth import get_user_model
    User = get_user_model()
    
    # Get filter parameters
    selected_period = request.GET.get('grading_period', 'overall')
    selected_section = request.GET.get('section', 'all')
    search_query = request.GET.get('search', '').strip()

    students_qs = User.objects.filter(is_staff=False)
    if selected_section != 'all':
        try:
            from students.models import Student
            student_ids = Student.objects.filter(section=selected_section).values_list('user_id', flat=True)
            students_qs = students_qs.filter(id__in=student_ids)
        except Exception:
            if hasattr(User, 'section'):
                students_qs = students_qs.filter(section=selected_section)
    if search_query:
        students_qs = students_qs.filter(
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(username__icontains=search_query)
        )
    all_students = students_qs
    
    # Get section choices robustly
    section_choices = []
    try:
        from students.models import SECTION_CHOICES
        section_choices = SECTION_CHOICES
    except ImportError:
        try:
            from accounts.models import SECTION_CHOICES
            section_choices = SECTION_CHOICES
        except ImportError:
            # Try to build choices from existing students if the static is missing
            section_set = set()
            try:
                from students.models import Student
                section_set = set(Student.objects.values_list('section', flat=True))
            except Exception:
                pass
            if not section_set:
                # Fallback: guess from User objects if they have 'section' field
                if hasattr(User, 'section'):
                    section_set = set(User.objects.exclude(section=None).values_list('section', flat=True))
            section_choices = [(sec, sec) for sec in section_set if sec]
    # Ensure at least "All Sections" option exists
    if section_choices and ("all", "All Sections") not in section_choices:
        section_choices = [("all", "All Sections")] + list(section_choices)
    elif not section_choices:
        section_choices = [("all", "All Sections")]

    # Use grading period choices from lessons model
    grading_periods = LESSON_GRADING_PERIOD_CHOICES

    total_users = User.objects.filter(is_superuser=False).count()
    new_users_week = User.objects.filter(
        date_joined__gte=timezone.now() - timedelta(days=7)
    ).count()
    active_users = User.objects.filter(
        last_login__gte=timezone.now() - timedelta(days=30)
    ).count()
    total_lessons = Lesson.objects.filter(is_active=True, is_archived=False).count()
    total_quizzes = Quiz.objects.filter(is_published=True, is_archived=False).count()
    total_exams = Exam.objects.count()

    recent_activities = ActivityLog.objects.all().order_by('-timestamp')[:10]

    period_stats = {}

    for period_value, period_label in GRADING_PERIODS:
        total_lessons_in_period = Lesson.objects.filter(grading_period=period_value, is_active=True, is_archived=False).count()

        quizzes_in_period = Quiz.objects.filter(grading_period=period_value, is_archived=False)
        quizzes_list = [
            {'id': q.id, 'title': q.title}
            for q in quizzes_in_period
        ]

        exams_in_period = Exam.objects.filter(grading_period=period_value)
        exams_list = [
            {'id': e.id, 'title': e.title}
            for e in exams_in_period
        ]

        quiz_rankings = []
        for quiz in quizzes_in_period:
            attempts = (QuizAttempt.objects.filter(
                quiz=quiz, completed=True)
                .select_related('user')
                .order_by('-score', 'user__username'))

            if has_grading_period_field(QuizAttempt):
                attempts = attempts.filter(grading_period=period_value)

            student_best = {}
            for att in attempts:
                stu_id = att.user.id
                score = att.score
                earned_points = getattr(att, "raw_points", None)
                total_points = getattr(att, "total_points", None)
                if total_points is None:
                    total_points = getattr(quiz, 'max_score', 100)
                if earned_points is None:
                    earned_points = score
                if stu_id not in student_best or score > student_best[stu_id]['score']:
                    student_best[stu_id] = {
                        'student_name': att.user.get_full_name() or att.user.username,
                        'score': score,
                        'earned_points': earned_points,
                        'total_points': total_points,
                        'passed': getattr(att, "passed", None) if hasattr(att, "passed") else None
                    }
            ranked = sorted(student_best.values(), key=lambda x: x['score'], reverse=True)
            rankings_rows = []
            last_score = None
            rank = 0
            for idx, row in enumerate(ranked, 1):
                if last_score is None or row['score'] < last_score:
                    rank = idx
                rankings_rows.append({
                    'rank': rank,
                    'student_name': row['student_name'],
                    'score': row['score'],
                    'earned_points': row['earned_points'],
                    'total_points': row['total_points'],
                    'passed': row['passed']
                })
                last_score = row['score']
            quiz_rankings.append({
                'quiz_title': quiz.title,
                'rankings': rankings_rows,
            })

        exam_rankings = []
        for exam in exams_in_period:
            attempts = (ExamAttempt.objects.filter(
                exam=exam, completed=True)
                .select_related('user')
                .order_by('-score', 'user__username'))

            if has_grading_period_field(ExamAttempt):
                attempts = attempts.filter(grading_period=period_value)

            student_best = {}
            for att in attempts:
                stu_id = att.user.id
                score = att.score
                earned_points = getattr(att, "raw_points", None)
                total_points = getattr(att, "total_points", None)
                if total_points is None:
                    total_points = getattr(exam, 'max_score', 100)
                if earned_points is None:
                    earned_points = score
                if stu_id not in student_best or score > student_best[stu_id]['score']:
                    student_best[stu_id] = {
                        'student_name': att.user.get_full_name() or att.user.username,
                        'score': score,
                        'earned_points': earned_points,
                        'total_points': total_points,
                        'passed': getattr(att, "passed", None) if hasattr(att, "passed") else None
                    }
            ranked = sorted(student_best.values(), key=lambda x: x['score'], reverse=True)
            rankings_rows = []
            last_score = None
            rank = 0
            for idx, row in enumerate(ranked, 1):
                if last_score is None or row['score'] < last_score:
                    rank = idx
                rankings_rows.append({
                    'rank': rank,
                    'student_name': row['student_name'],
                    'score': row['score'],
                    'earned_points': row['earned_points'],
                    'total_points': row['total_points'],
                    'passed': row['passed']
                })
                last_score = row['score']
            exam_rankings.append({
                'exam_title': exam.title,
                'rankings': rankings_rows,
            })

        period_stats[period_value] = {
            'label': period_label,
            'total_lessons': total_lessons_in_period,
            'quizzes': quizzes_list,
            'exams': exams_list,
            'quiz_rankings': quiz_rankings,
            'exam_rankings': exam_rankings,
        }

    from django.db.models import Max, Avg
    quiz_stats = QuizAttempt.objects.filter(completed=True).aggregate(
        avg_score=Avg('score'),
        high_score=Max('score')
    )

    all_exam_attempts = ExamAttempt.objects.filter(completed=True)
    total_exam_attempts = all_exam_attempts.count()
    passing_exam_attempts = all_exam_attempts.filter(passed=True).count()

    exam_stats = ExamAttempt.objects.filter(completed=True).aggregate(
        avg_score=Avg('score')
    )

    if total_exam_attempts > 0:
        exam_stats['pass_rate'] = passing_exam_attempts / total_exam_attempts
    else:
        exam_stats['pass_rate'] = 0

    # Use the filtered students from above
    student_totals = []
    for student in all_students:
        quiz_qs = QuizAttempt.objects.filter(
            user=student,
            completed=True,
            quiz__is_archived=False 
        )
        exam_qs = ExamAttempt.objects.filter(user=student, completed=True)
        stu_quiz_points = quiz_qs.aggregate(s=Sum('raw_points'))['s']
        if stu_quiz_points is None:
            stu_quiz_points = quiz_qs.aggregate(s=Sum('score'))['s'] or 0
        stu_exam_points = exam_qs.aggregate(s=Sum('raw_points'))['s']
        if stu_exam_points is None:
            stu_exam_points = exam_qs.aggregate(s=Sum('score'))['s'] or 0
        student_totals.append({
            'user_id': student.id,
            'full_name': f"{student.last_name}, {student.first_name}".strip() or student.username,
            'total_points': (stu_quiz_points or 0) + (stu_exam_points or 0),
            'quiz_points': stu_quiz_points or 0,
            'exam_points': stu_exam_points or 0,
        })
    sorted_students = sorted(student_totals, key=lambda x: x['total_points'], reverse=True)
    overall_rankings = []
    for idx, stu in enumerate(sorted_students, 1):
        overall_rankings.append({
            'rank': idx,
            'full_name': stu['full_name'],
            'quiz_points': stu['quiz_points'],
            'exam_points': stu['exam_points'],
            'total_points': stu['total_points'],
        })

    # Calculate student grades with null handling
    students_grades = []
    try:
        from students.models import Student
        has_student_model = True
    except ImportError:
        has_student_model = False

    grading_period_keys = [p[0] for p in grading_periods]

    if selected_period == 'overall':
        # For each student, get grades for each period, then average
        for student in all_students:
            # Section info
            section_name = None
            if has_student_model:
                try:
                    student_profile = Student.objects.get(user=student)
                    section_name = getattr(student_profile, "section", None)
                    if callable(section_name):
                        section_name = section_name()
                    elif hasattr(section_name, "name"):
                        section_name = section_name.name
                    elif isinstance(section_name, str):
                        section_name = section_name
                    else:
                        section_name = str(section_name or "N/A")
                except Student.DoesNotExist:
                    section_name = "N/A"
            elif hasattr(student, "section"):
                section_name = getattr(student, "section", "N/A")

            period_grades = {}
            period_count = 0
            period_sum = 0
            for period in grading_period_keys:
                quiz_filter = dict(user=student, completed=True, quiz__is_archived=False, grading_period=period)
                exam_filter = dict(user=student, completed=True, grading_period=period)
                quiz_avg = QuizAttempt.objects.filter(**quiz_filter).aggregate(avg=Avg('score'))['avg']
                exam_avg = ExamAttempt.objects.filter(**exam_filter).aggregate(avg=Avg('score'))['avg']
                if quiz_avg is not None and exam_avg is not None:
                    grade = quiz_avg * 0.4 + exam_avg * 0.6
                    period_grades[period] = grade
                    period_sum += grade
                    period_count += 1
                else:
                    period_grades[period] = None
            overall_grade = (period_sum / period_count) if period_count > 0 else None
            # GPA calculation (same as before)
            gpa = None
            if overall_grade is not None:
                if overall_grade >= 99:
                    gpa = 1.00
                elif overall_grade >= 96:
                    gpa = 1.25
                elif overall_grade >= 93:
                    gpa = 1.50
                elif overall_grade >= 90:
                    gpa = 1.75
                elif overall_grade >= 87:
                    gpa = 2.00
                elif overall_grade >= 84:
                    gpa = 2.25
                elif overall_grade >= 81:
                    gpa = 2.50
                elif overall_grade >= 78:
                    gpa = 2.75
                elif overall_grade >= 75:
                    gpa = 3.00
                else:
                    gpa = 5.00
            students_grades.append({
                'section': section_name or "N/A",
                'full_name': f"{student.last_name}, {student.first_name}".strip() or student.username,
                'prelim': period_grades.get('prelim'),
                'midterm': period_grades.get('midterm'),
                'prefinal': period_grades.get('prefinal'),
                'final': period_grades.get('final'),
                'overall': overall_grade,
                'gpa': gpa,
            })
    else:
        # For a specific period, get weighted scores for written works, performance task, exams
        # Written Works (20%)
        written_works_quizzes = list(Quiz.objects.filter(grading_period=selected_period, is_archived=False).order_by('id'))
        for student in all_students:
            # Section info
            section_name = None
            if has_student_model:
                try:
                    student_profile = Student.objects.get(user=student)
                    section_name = getattr(student_profile, "section", None)
                    if callable(section_name):
                        section_name = section_name()
                    elif hasattr(section_name, "name"):
                        section_name = section_name.name
                    elif isinstance(section_name, str):
                        section_name = section_name
                    else:
                        section_name = str(section_name or "N/A")
                except Student.DoesNotExist:
                    section_name = "N/A"
            elif hasattr(student, "section"):
                section_name = getattr(student, "section", "N/A")
            # Written Works
            quiz_scores = []
            student_has_attempt = False
            for quiz in written_works_quizzes:
                attempt = QuizAttempt.objects.filter(user=student, quiz=quiz, completed=True).order_by('-score').first()
                if attempt:
                    student_has_attempt = True
                    raw_points = getattr(attempt, 'raw_points', None)
                    total_points = getattr(attempt, 'total_points', None)
                    if raw_points is None:
                        raw_points = attempt.score
                    if total_points is None:
                        total_points = getattr(quiz, 'max_score', 100)
                    percent = round((raw_points / total_points) * 100, 2) if total_points else 0.0
                    quiz_scores.append(percent)
                else:
                    quiz_scores.append(0.0)
            if written_works_quizzes:
                if student_has_attempt:
                    ww_percentage = round(sum(quiz_scores) / len(written_works_quizzes), 2)
                    ww_weighted = round(ww_percentage * 0.2, 2)
                else:
                    ww_percentage = None
                    ww_weighted = None
            else:
                ww_percentage = None
                ww_weighted = None
            # Performance Task (50%)
            if hasattr(Attendance, 'grading_period'):
                attendance_qs = Attendance.objects.filter(user=student, grading_period=selected_period)
            else:
                attendance_qs = Attendance.objects.filter(user=student)
            days_present = attendance_qs.aggregate(total=Sum('days_present'))['total'] or 0
            total_days = attendance_qs.aggregate(total=Sum('total_days'))['total'] or 0
            if total_days > 0:
                attendance_percent = round((days_present / total_days) * 100, 2)
            else:
                attendance_percent = 0.0
            pt_weighted = round(attendance_percent * 0.5, 2) if total_days > 0 else 0.0
            # Examinations (30%)
            exam_attempt = ExamAttempt.objects.filter(user=student, completed=True, grading_period=selected_period).order_by('-id').first()
            if exam_attempt:
                user_score = getattr(exam_attempt, 'raw_points', None)
                total_items = getattr(exam_attempt, 'total_points', None)
                if user_score is None:
                    user_score = getattr(exam_attempt, 'score', None)
                if total_items is None:
                    total_items = 100
                exam_percent = round((user_score / total_items) * 100, 2) if user_score is not None and total_items else None
                exam_weighted = round(exam_percent * 0.3, 2) if exam_percent is not None else None
            else:
                exam_percent = None
                exam_weighted = None
            # Grade (sum)
            if ww_weighted is not None and pt_weighted is not None and exam_weighted is not None:
                grade_sum = round(ww_weighted + pt_weighted + exam_weighted, 2)
            else:
                grade_sum = None
            # GPA calculation (same as before)
            gpa = None
            if grade_sum is not None:
                if grade_sum >= 99:
                    gpa = 1.00
                elif grade_sum >= 96:
                    gpa = 1.25
                elif grade_sum >= 93:
                    gpa = 1.50
                elif grade_sum >= 90:
                    gpa = 1.75
                elif grade_sum >= 87:
                    gpa = 2.00
                elif grade_sum >= 84:
                    gpa = 2.25
                elif grade_sum >= 81:
                    gpa = 2.50
                elif grade_sum >= 78:
                    gpa = 2.75
                elif grade_sum >= 75:
                    gpa = 3.00
                else:
                    gpa = 5.00
            students_grades.append({
                'section': section_name or "N/A",
                'full_name': f"{student.last_name}, {student.first_name}".strip() or student.username,
                'written_works': ww_weighted,
                'performance_task': pt_weighted,
                'examinations': exam_weighted,
                'grade': grade_sum,
                'gpa': gpa,
            })
    # --- SORTING FOR STUDENTS GRADES TABLE ---
    sort = request.GET.get('sort')
    dir = request.GET.get('dir', 'asc')
    if sort:
        reverse = (dir == 'desc')
        def sort_key(x):
            v = x.get(sort)
            # For None values, sort them last
            if v is None:
                return float('-inf') if reverse else float('inf')
            return v
        students_grades.sort(key=sort_key, reverse=reverse)

    # Written Works (20%) Card logic
    selected_ww_period = request.GET.get('ww_grading_period', 'prelim')
    selected_ww_section = request.GET.get('ww_section', 'all')
    ww_search_query = request.GET.get('ww_search', '').strip()

    ww_students_qs = User.objects.filter(is_staff=False)
    if selected_ww_section != 'all':
        try:
            from students.models import Student
            student_ids = Student.objects.filter(section=selected_ww_section).values_list('user_id', flat=True)
            ww_students_qs = ww_students_qs.filter(id__in=student_ids)
        except Exception:
            if hasattr(User, 'section'):
                ww_students_qs = ww_students_qs.filter(section=selected_ww_section)
    if ww_search_query:
        ww_students_qs = ww_students_qs.filter(
            Q(first_name__icontains=ww_search_query) |
            Q(last_name__icontains=ww_search_query) |
            Q(username__icontains=ww_search_query)
        )

    # Get all quizzes for the selected grading period
    written_works_quizzes = list(Quiz.objects.filter(grading_period=selected_ww_period, is_archived=False).order_by('id'))

    written_works_grades = []
    try:
        from students.models import Student
        has_student_model = True
    except ImportError:
        has_student_model = False
    for student in ww_students_qs:
        # Section info
        section_name = None
        if has_student_model:
            try:
                student_profile = Student.objects.get(user=student)
                section_name = getattr(student_profile, "section", None)
                if callable(section_name):
                    section_name = section_name()
                elif hasattr(section_name, "name"):
                    section_name = section_name.name
                elif isinstance(section_name, str):
                    section_name = section_name
                else:
                    section_name = str(section_name or "N/A")
            except Student.DoesNotExist:
                section_name = "N/A"
        elif hasattr(student, "section"):
            section_name = getattr(student, "section", "N/A")

        quiz_scores = []
        quiz_score_map = {}
        student_has_attempt = False
        for quiz in written_works_quizzes:
            attempt = QuizAttempt.objects.filter(user=student, quiz=quiz, completed=True).order_by('-score').first()
            if attempt:
                student_has_attempt = True
                raw_points = getattr(attempt, 'raw_points', None)
                total_points = getattr(attempt, 'total_points', None)
                if raw_points is None:
                    raw_points = attempt.score
                if total_points is None:
                    total_points = getattr(quiz, 'max_score', 100)
                percent = round((raw_points / total_points) * 100, 2) if total_points else 0.0
                remarks = "Passed" if percent is not None and percent >= 75 else "Failed"
                quiz_scores.append(percent)
                quiz_score_map[quiz.id] = {
                    'raw_points': raw_points,
                    'total_points': total_points,
                    'percent': percent,
                    'remarks': remarks,
                }
            else:
                # No attempt: treat as 0%
                quiz_scores.append(0.0)
                quiz_score_map[quiz.id] = {
                    'raw_points': None,
                    'total_points': None,
                    'percent': 0.0,
                    'remarks': "Failed",
                }

        # Percentage score is the average of all quiz percentages (including 0 for missing)
        if written_works_quizzes:
            if student_has_attempt:
                percentage_score = round(sum(quiz_scores) / len(written_works_quizzes), 2)
                weighted_score = round(percentage_score * 0.2, 2)
            else:
                percentage_score = None
                weighted_score = None
        else:
            percentage_score = None
            weighted_score = None
        # Compute overall remarks:
        # - '--' if no quizzes in period
        # - '--' if student has no attempts for any quiz in period
        # - 'Passed' if all quizzes are Passed
        # - 'Failed' otherwise
        if not written_works_quizzes:
            overall_remarks = '--'
        elif not student_has_attempt:
            overall_remarks = '--'
        else:
            all_passed = all(q['remarks'] == 'Passed' for q in quiz_score_map.values() if q is not None)
            overall_remarks = 'Passed' if all_passed else 'Failed'
        written_works_grades.append({
            'section': section_name or "N/A",
            'full_name': f"{student.last_name}, {student.first_name}".strip() or student.username,
            'quiz_scores': quiz_score_map,
            'percentage_score': percentage_score,
            'weighted_score': weighted_score,
            'overall_remarks': overall_remarks,
        })

    # --- SORTING FOR WRITTEN WORKS TABLE ---
    ww_sort = request.GET.get('sort')
    ww_dir = request.GET.get('dir', 'asc')
    if ww_sort:
        reverse = (ww_dir == 'desc')
        def ww_sort_key(x):
            # Quiz columns: quiz_{id}
            if ww_sort.startswith('quiz_'):
                quiz_id = None
                try:
                    quiz_id = int(ww_sort.split('_')[1])
                except Exception:
                    return 0
                quiz_score = x['quiz_scores'].get(quiz_id, {}).get('percent', 0)
                return quiz_score if quiz_score is not None else (float('-inf') if reverse else float('inf'))
            # Other columns
            v = x.get(ww_sort)
            if v is None:
                return float('-inf') if reverse else float('inf')
            return v
        written_works_grades.sort(key=ww_sort_key, reverse=reverse)

    # Exam Grades Card logic
    selected_exam_period = request.GET.get('exam_grading_period', 'prelim')
    selected_exam_section = request.GET.get('exam_section', 'all')
    exam_search_query = request.GET.get('exam_search', '').strip()

    exam_students_qs = User.objects.filter(is_staff=False)
    if selected_exam_section != 'all':
        try:
            from students.models import Student
            student_ids = Student.objects.filter(section=selected_exam_section).values_list('user_id', flat=True)
            exam_students_qs = exam_students_qs.filter(id__in=student_ids)
        except Exception:
            if hasattr(User, 'section'):
                exam_students_qs = exam_students_qs.filter(section=selected_exam_section)
    if exam_search_query:
        exam_students_qs = exam_students_qs.filter(
            Q(first_name__icontains=exam_search_query) |
            Q(last_name__icontains=exam_search_query) |
            Q(username__icontains=exam_search_query)
        )
    exam_grades = []
    try:
        from students.models import Student
        has_student_model = True
    except ImportError:
        has_student_model = False
    for student in exam_students_qs:
        # Section info
        section_name = None
        if has_student_model:
            try:
                student_profile = Student.objects.get(user=student)
                section_name = getattr(student_profile, "section", None)
                if callable(section_name):
                    section_name = section_name()
                elif hasattr(section_name, "name"):
                    section_name = section_name.name
                elif isinstance(section_name, str):
                    section_name = section_name
                else:
                    section_name = str(section_name or "N/A")
            except Student.DoesNotExist:
                section_name = "N/A"
        elif hasattr(student, "section"):
            section_name = getattr(student, "section", "N/A")

        # Only consider the selected grading period
        grading_period = selected_exam_period
        if grading_period == 'overall':
            grading_period = None

        # Get latest ExamAttempt for this student and grading period
        exam_attempt = None
        if grading_period:
            exam_attempt = ExamAttempt.objects.filter(
                user=student, completed=True, grading_period=grading_period
            ).order_by('-id').first()
        else:
            exam_attempt = ExamAttempt.objects.filter(
                user=student, completed=True
            ).order_by('-id').first()

        if exam_attempt:
            user_score = getattr(exam_attempt, 'raw_points', None)
            total_items = getattr(exam_attempt, 'total_points', None)
            if user_score is None:
                user_score = getattr(exam_attempt, 'score', None)
            if total_items is None:
                total_items = 100
            percent_score = round((user_score / total_items) * 100, 2) if user_score is not None and total_items else None
            weighted_score = round(percent_score * 0.3, 2) if percent_score is not None else None
            passed = getattr(exam_attempt, 'passed', None)
            remarks = None
            if passed is not None:
                remarks = 'Passed' if passed else 'Failed'
            elif percent_score is not None:
                remarks = 'Passed' if percent_score >= 75 else 'Failed'
            else:
                remarks = '--'
            score_display = f"{int(user_score) if user_score is not None else '--'}/{int(total_items) if total_items is not None else '--'}"
        else:
            user_score = None
            total_items = None
            percent_score = None
            weighted_score = None
            remarks = '--'
            score_display = '--/--'

        exam_grades.append({
            'section': section_name or "N/A",
            'full_name': f"{student.last_name}, {student.first_name}".strip() or student.username,
            'score_display': score_display,
            'percent_score': percent_score,
            'weighted_score': weighted_score,
            'remarks': remarks,
        })

    # --- SORTING FOR EXAM GRADES TABLE ---
    exam_sort = request.GET.get('sort')
    exam_dir = request.GET.get('dir', 'asc')
    if exam_sort:
        reverse = (exam_dir == 'desc')
        def exam_sort_key(x):
            val = x.get(exam_sort)
            if val is None:
                return float('-inf') if reverse else float('inf')
            return val
        exam_grades.sort(key=exam_sort_key, reverse=reverse)

    # --- PERFORMANCE TASKS TABLE LOGIC ---
    selected_pt_period = request.GET.get('pt_grading_period', 'prelim')
    selected_pt_section = request.GET.get('pt_section', 'all')
    pt_search_query = request.GET.get('pt_search', '').strip()

    pt_students_qs = User.objects.filter(is_staff=False)
    if selected_pt_section != 'all':
        try:
            from students.models import Student
            student_ids = Student.objects.filter(section=selected_pt_section).values_list('user_id', flat=True)
            pt_students_qs = pt_students_qs.filter(id__in=student_ids)
        except Exception:
            if hasattr(User, 'section'):
                pt_students_qs = pt_students_qs.filter(section=selected_pt_section)
    if pt_search_query:
        pt_students_qs = pt_students_qs.filter(
            Q(first_name__icontains=pt_search_query) |
            Q(last_name__icontains=pt_search_query) |
            Q(username__icontains=pt_search_query)
        )

    performance_tasks = []
    try:
        from students.models import Student
        has_student_model = True
    except ImportError:
        has_student_model = False
    for student in pt_students_qs:
        # Section info
        section_name = None
        if has_student_model:
            try:
                student_profile = Student.objects.get(user=student)
                section_name = getattr(student_profile, "section", None)
                if callable(section_name):
                    section_name = section_name()
                elif hasattr(section_name, "name"):
                    section_name = section_name.name
                elif isinstance(section_name, str):
                    section_name = section_name
                else:
                    section_name = str(section_name or "N/A")
            except Student.DoesNotExist:
                section_name = "N/A"
        elif hasattr(student, "section"):
            section_name = getattr(student, "section", "N/A")

        # Attendance: sum of days_present and total_days in selected period
        if hasattr(Attendance, 'grading_period'):
            attendance_qs = Attendance.objects.filter(user=student, grading_period=selected_pt_period)
        else:
            attendance_qs = Attendance.objects.filter(user=student)
        days_present = attendance_qs.aggregate(total=Sum('days_present'))['total'] or 0
        total_days = attendance_qs.aggregate(total=Sum('total_days'))['total'] or 0
        if total_days > 0:
            attendance_percent = round((days_present / total_days) * 100, 2)
        else:
            attendance_percent = 0.0

        # Performance Task: percent_score and weighted_score (attendance-based)
        percent_score = attendance_percent
        weighted_score = round(percent_score * 0.5, 2) if total_days > 0 else 0.0

        performance_tasks.append({
            'user_id': student.id,
            'section': section_name or "N/A",
            'full_name': f"{student.last_name}, {student.first_name}".strip() or student.username,
            'attendance': days_present,  # for backward compatibility
            'days_present': days_present,
            'total_days': total_days,
            'attendance_percent': attendance_percent,
            'percent_score': percent_score,
            'weighted_score': weighted_score,
        })

    context = {
        'total_users': total_users,
        'new_users_week': new_users_week,
        'active_users': active_users,
        'total_lessons': total_lessons,
        'total_quizzes': total_quizzes,
        'total_exams': total_exams,
        'recent_activities': recent_activities,
        'quiz_stats': quiz_stats,
        'exam_stats': exam_stats,
        'grading_periods': grading_periods,
        'section_choices': section_choices,
        'selected_period': selected_period,
        'selected_section': selected_section,
        'search_query': search_query,
        'period_stats': period_stats,
        'overall_rankings': overall_rankings,
        'students_grades': students_grades,
        'exam_grades': exam_grades,
        'sort': sort,
        'dir': dir,
        'exam_sort': exam_sort,
        'exam_dir': exam_dir,
        'selected_exam_period': selected_exam_period,
        'selected_exam_section': selected_exam_section,
        'exam_search_query': exam_search_query,
        'written_works_grades': written_works_grades,
        'written_works_quizzes': written_works_quizzes,
        'selected_ww_period': selected_ww_period,
        'selected_ww_section': selected_ww_section,
        'ww_search_query': ww_search_query,
        'performance_tasks': performance_tasks,
        'selected_pt_period': selected_pt_period,
        'selected_pt_section': selected_pt_section,
        'pt_search_query': pt_search_query,
    }

    from django.template.loader import render_to_string
    # AJAX for Students Grades table
    if request.GET.get('ajax') == '1' and request.GET.get('table') == 'students_grades':
        html = render_to_string('dashboard/_students_grades_table_body.html', context, request=request)
        return HttpResponse(html)
    # AJAX for Exam Grades table
    elif request.GET.get('ajax') == '1' and request.GET.get('table') == 'exam_grades':
        html = render_to_string('dashboard/_exam_grades_table_body.html', context, request=request)
        return HttpResponse(html)
    # AJAX for Written Works table
    elif request.GET.get('ajax') == '1' and request.GET.get('table') == 'written_works':
        html = render_to_string('dashboard/_written_works_table_body.html', context, request=request)
        return HttpResponse(html)
    # AJAX for Performance Task table
    elif request.GET.get('ajax') == '1' and request.GET.get('table') == 'performance_task':
        html = render_to_string('dashboard/_performance_task_table_body.html', context, request=request)
        return HttpResponse(html)
    # Fallback: any AJAX
    elif request.GET.get('ajax') == '1' or request.headers.get('x-requested-with') == 'XMLHttpRequest':
        html = render_to_string('dashboard/_students_grades_table_body.html', context, request=request)
        return HttpResponse(html)
    else:
        return render(request, 'dashboard/admin.html', context)



@login_required
def performance_task_view(request):
    """
    View for displaying performance tasks, including attendance and student progress.
    """
    attendance_records = Attendance.objects.select_related('user').order_by('-date')
    student_progress = StudentProgress.objects.select_related('user').all()
    context = {
        'attendance_records': attendance_records,
        'student_progress': student_progress,
    }
    return render(request, 'dashboard/performance_task.html', context)

@login_required
def download_rankings_docx(request, period_value=None):
    TEMPLATE_PATH = os.path.join(settings.BASE_DIR, 'media', 'templates', 'RANKING_TEMPLATE.docx')


    from django.contrib.auth import get_user_model
    User = get_user_model()
    
    # Get filter parameters
    selected_section = request.GET.get('section', 'all')
    search_query = request.GET.get('search', '').strip()
    
    # Initial queryset: all non-staff students
    students_qs = User.objects.filter(is_staff=False)
    
    # Apply section filter if available
    if selected_section != 'all':
        # Attempt to join with students or from User if 'section' field
        try:
            from students.models import Student
            student_ids = Student.objects.filter(section=selected_section).values_list('user_id', flat=True)
            students_qs = students_qs.filter(id__in=student_ids)
        except Exception:
            # Fallback: check User.section if available
            if hasattr(User, 'section'):
                students_qs = students_qs.filter(section=selected_section)
    
    # Apply search filter
    if search_query:
        students_qs = students_qs.filter(
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(username__icontains=search_query)
        )
    
    # Define all_students as the filtered queryset
    all_students = students_qs


    tables_to_include = request.GET.getlist('tables[]')
    if not tables_to_include:
        tables_to_include = ['overall', 'quizzes', 'exams']

    period_stats = {}
    for period_value_iter, period_label in GRADING_PERIODS:
        total_lessons_in_period = Lesson.objects.filter(grading_period=period_value_iter, is_active=True, is_archived=False).count()

        quizzes_in_period = Quiz.objects.filter(grading_period=period_value_iter, is_archived=False)
        quizzes_list = [
            {'id': q.id, 'title': q.title}
            for q in quizzes_in_period
        ]

        exams_in_period = Exam.objects.filter(grading_period=period_value_iter)
        exams_list = [
            {'id': e.id, 'title': e.title}
            for e in exams_in_period
        ]

        quiz_rankings = []
        for quiz in quizzes_in_period:
            attempts = (QuizAttempt.objects.filter(
                quiz=quiz, completed=True)
                .select_related('user')
                .order_by('-score', 'user__username'))
            if has_grading_period_field(QuizAttempt):
                attempts = attempts.filter(grading_period=period_value_iter)

            student_best = {}
            for att in attempts:
                stu_id = att.user.id
                score = att.score
                earned_points = getattr(att, "raw_points", None)
                total_points = getattr(att, "total_points", None)
                if total_points is None:
                    total_points = getattr(quiz, 'max_score', 100)
                if earned_points is None:
                    earned_points = score
                if stu_id not in student_best or score > student_best[stu_id]['score']:
                    student_best[stu_id] = {
                        'student_name': att.user.get_full_name() or att.user.username,
                        'score': score,
                        'earned_points': earned_points,
                        'total_points': total_points,
                        'passed': getattr(att, "passed", None) if hasattr(att, "passed") else None
                    }
            ranked = sorted(student_best.values(), key=lambda x: x['score'], reverse=True)
            rankings_rows = []
            last_score = None
            rank = 0
            for idx, row in enumerate(ranked, 1):
                if last_score is None or row['score'] < last_score:
                    rank = idx
                rankings_rows.append({
                    'rank': rank,
                    'student_name': row['student_name'],
                    'score': row['score'],
                    'earned_points': row['earned_points'],
                    'total_points': row['total_points'],
                    'passed': row['passed']
                })
                last_score = row['score']
            quiz_rankings.append({
                'quiz_title': quiz.title,
                'rankings': rankings_rows,
            })

        exam_rankings = []
        for exam in exams_in_period:
            attempts = (ExamAttempt.objects.filter(
                exam=exam, completed=True)
                .select_related('user')
                .order_by('-score', 'user__username'))
            if has_grading_period_field(ExamAttempt):
                attempts = attempts.filter(grading_period=period_value_iter)

            student_best = {}
            for att in attempts:
                stu_id = att.user.id
                score = att.score
                earned_points = getattr(att, "raw_points", None)
                total_points = getattr(att, "total_points", None)
                if total_points is None:
                    total_points = getattr(exam, 'max_score', 100)
                if earned_points is None:
                    earned_points = score
                if stu_id not in student_best or score > student_best[stu_id]['score']:
                    student_best[stu_id] = {
                        'student_name': att.user.get_full_name() or att.user.username,
                        'score': score,
                        'earned_points': earned_points,
                        'total_points': total_points,
                        'passed': getattr(att, "passed", None) if hasattr(att, "passed") else None
                    }
            ranked = sorted(student_best.values(), key=lambda x: x['score'], reverse=True)
            rankings_rows = []
            last_score = None
            rank = 0
            for idx, row in enumerate(ranked, 1):
                if last_score is None or row['score'] < last_score:
                    rank = idx
                rankings_rows.append({
                    'rank': rank,
                    'student_name': row['student_name'],
                    'score': row['score'],
                    'earned_points': row['earned_points'],
                    'total_points': row['total_points'],
                    'passed': row['passed']
                })
                last_score = row['score']
            exam_rankings.append({
                'exam_title': exam.title,
                'rankings': rankings_rows,
            })

        period_stats[period_value_iter] = {
            'label': period_label,
            'total_lessons': total_lessons_in_period,
            'quizzes': quizzes_list,
            'exams': exams_list,
            'quiz_rankings': quiz_rankings,
            'exam_rankings': exam_rankings,
        }

    # Get filter parameters
    selected_section = request.GET.get('section', 'all')
    search_query = request.GET.get('search', '').strip()
    
    # Use the filtered students from above
    student_totals = []
    for student in all_students:
        quiz_qs = QuizAttempt.objects.filter(
            user=student,
            completed=True,
            quiz__is_archived=False
        )
        exam_qs = ExamAttempt.objects.filter(user=student, completed=True)
        stu_quiz_points = quiz_qs.aggregate(s=Sum('raw_points'))['s']
        if stu_quiz_points is None:
            stu_quiz_points = quiz_qs.aggregate(s=Sum('score'))['s'] or 0
        stu_exam_points = exam_qs.aggregate(s=Sum('raw_points'))['s']
        if stu_exam_points is None:
            stu_exam_points = exam_qs.aggregate(s=Sum('score'))['s'] or 0
        student_totals.append({
            'user_id': student.id,
            'full_name': f"{student.first_name} {student.last_name}".strip() or student.username,
            'total_points': (stu_quiz_points or 0) + (stu_exam_points or 0),
            'quiz_points': stu_quiz_points or 0,
            'exam_points': stu_exam_points or 0,
        })
    sorted_students = sorted(student_totals, key=lambda x: x['total_points'], reverse=True)
    overall_rankings = []
    for idx, stu in enumerate(sorted_students, 1):
        overall_rankings.append({
            'rank': idx,
            'full_name': stu['full_name'],
            'quiz_points': stu['quiz_points'],
            'exam_points': stu['exam_points'],
            'total_points': stu['total_points'],
        })

    def apply_modern_table_style(table, header_color=RGBColor(0xE1, 0xE0, 0xE0)):
        color_hex = str(header_color)
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    if row_idx == 0:
                        for run in paragraph.runs:
                            run.bold = True
                        tc_pr = cell._tc.get_or_add_tcPr()
                        shd_xml = (
                            f'<w:shd {nsdecls("w")}'
                            f' w:val="clear"'
                            f' w:color="auto"'
                            f' w:fill="{color_hex}"/>'
                        )
                        tc_pr.append(parse_xml(shd_xml))
                    if col_idx != 1:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def int_format(value):
        try:
            return str(int(round(float(value))))
        except Exception:
            return str(value)


    def add_docx_heading(doc, text, level=1):
        style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
        default_styles = ["Heading 1", "Heading 2", "Heading 3", "Title", "Subtitle"]
        style_name = style_map.get(level, "Heading 1")
        for try_style in [style_name] + default_styles:
            if try_style in [s.name for s in doc.styles]:
                return doc.add_paragraph(text, style=try_style)
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.bold = True
            if level == 1:
                run.font.size = docx.shared.Pt(16)
            elif level == 2:
                run.font.size = docx.shared.Pt(14)
            elif level == 3:
                run.font.size = docx.shared.Pt(12)
        return p

    doc = Document(TEMPLATE_PATH)
    add_docx_heading(doc, 'OFFICIAL STUDENT RANKINGS', level=1)
    doc.add_paragraph(f"Generated on: {datetime.date.today().strftime('%B %d, %Y')}")
    doc.add_paragraph()  

    if 'overall' in tables_to_include:
        add_docx_heading(doc, 'Overall Student Rankings', level=2)
        doc.add_paragraph('Based on cumulative quiz and exam points', style='Intense Quote' if 'Intense Quote' in [s.name for s in doc.styles] else None)
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        header_texts = ['Rank', 'Student Name', 'Quiz Score', 'Exam Score', 'Total Score']
        for i, text in enumerate(header_texts):
            hdr_cells[i].text = text
        for stu in overall_rankings:
            row_cells = table.add_row().cells
            row_data = [
                f"#{stu['rank']}",
                stu['full_name'],
                int_format(stu['quiz_points']),
                int_format(stu['exam_points']),
                int_format(stu['total_points'])
            ]
            for i, data in enumerate(row_data):
                row_cells[i].text = data
        apply_modern_table_style(table)
        doc.add_page_break()


    periods_to_show = [(p_val, p_label) for p_val, p_label in GRADING_PERIODS
                      if period_value is None or p_val == period_value]
    for p_val, p_label in periods_to_show:
        stats = period_stats[p_val]
        any_period_tables = False

        if 'quizzes' in tables_to_include and stats.get('quiz_rankings'):
            any_period_tables = True
        if 'exams' in tables_to_include and stats.get('exam_rankings'):
            any_period_tables = True

        if not any_period_tables:
            continue

        add_docx_heading(doc, f'{p_label} Period Rankings', level=1)
        doc.add_paragraph()  
        if 'quizzes' in tables_to_include and stats.get('quiz_rankings'):
            add_docx_heading(doc, 'Quiz Rankings', level=2)
            for quiz in stats['quiz_rankings']:
                add_docx_heading(doc, quiz['quiz_title'], level=3)
                quiz_table = doc.add_table(rows=1, cols=5)
                quiz_header = quiz_table.rows[0].cells
                header_texts = ['Rank', 'Student', 'Percentage', 'Score', 'Status']
                for i, text in enumerate(header_texts):
                    quiz_header[i].text = text
                for rank in quiz['rankings']:
                    r = quiz_table.add_row().cells
                    row_data = [
                        f"#{rank['rank']}",
                        rank['student_name'],
                        int_format(rank['score']),
                        f"{int_format(rank['earned_points'])} / {int_format(rank['total_points'])}",
                        "Passed" if rank['passed'] else "Failed"
                    ]
                    for i, data in enumerate(row_data):
                        r[i].text = data
                apply_modern_table_style(quiz_table)
                doc.add_paragraph()

        if 'exams' in tables_to_include and stats.get('exam_rankings'):
            add_docx_heading(doc, 'Exam Rankings', level=2)
            for exam in stats['exam_rankings']:
                add_docx_heading(doc, exam['exam_title'], level=3)
                exam_table = doc.add_table(rows=1, cols=5)
                exam_header = exam_table.rows[0].cells
                header_texts = ['Rank', 'Student', 'Percentage', 'Score', 'Status']
                for i, text in enumerate(header_texts):
                    exam_header[i].text = text
                for rank in exam['rankings']:
                    r = exam_table.add_row().cells
                    row_data = [
                        f"#{rank['rank']}",
                        rank['student_name'],
                        int_format(rank['score']),
                        f"{int_format(rank['earned_points'])} / {int_format(rank['total_points'])}",
                        "Passed" if rank['passed'] else "Failed"
                    ]
                    for i, data in enumerate(row_data):
                        r[i].text = data
                apply_modern_table_style(exam_table)
                doc.add_paragraph()
        if p_val != periods_to_show[-1][0]:
            doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    period_suffix = f"_{period_value}" if period_value else ""
    filename = f"Student_Rankings_Report{period_suffix}_{datetime.date.today().isoformat()}.docx"
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@require_POST
def edit_total_days(request):
    grading_period = request.POST.get('grading_period')
    total_days = request.POST.get('total_days')
    try:
        total_days = int(total_days)
    except (TypeError, ValueError):
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            from django.http import JsonResponse
            return JsonResponse({'success': False, 'message': 'Invalid number of days.'}, status=400)
        messages.error(request, "Invalid number of days.")
        return redirect('admin_dashboard')

    from .models import Attendance
    from django.contrib.auth import get_user_model
    User = get_user_model()

    # Only filter if grading_period is not 'overall' and not empty
    if grading_period and grading_period != 'overall':
        students = User.objects.filter(is_staff=False)
        # Ensure every student has an Attendance record for this period
        for student in students:
            Attendance.objects.get_or_create(user=student, grading_period=grading_period)
        attendance_qs = Attendance.objects.filter(grading_period=grading_period)
    else:
        students = User.objects.filter(is_staff=False)
        # Ensure every student has an Attendance record for every period
        from lessons.models import GRADING_PERIOD_CHOICES
        for student in students:
            for period, _ in GRADING_PERIOD_CHOICES:
                Attendance.objects.get_or_create(user=student, grading_period=period)
        attendance_qs = Attendance.objects.all()
    updated = attendance_qs.update(total_days=total_days)
    msg = f"Total days updated for {updated} students."
    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        from django.http import JsonResponse
        return JsonResponse({'success': True, 'message': msg})
    messages.success(request, msg)
    return redirect('admin_dashboard')



@require_POST
def edit_days_present(request):
    user_id = request.POST.get('user_id')
    grading_period = request.POST.get('grading_period')
    days_present = request.POST.get('days_present')

    # Validate input
    if not user_id or not grading_period or days_present is None:
        return JsonResponse({'success': False, 'message': 'Missing required fields.'}, status=400)
    try:
        days_present = int(days_present)
        if days_present < 0:
            raise ValueError
    except ValueError:
        return JsonResponse({'success': False, 'message': 'Invalid days present.'}, status=400)

    User = get_user_model()
    try:
        user = User.objects.get(pk=user_id)
    except User.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'User not found.'}, status=404)

    # Update or create Attendance record
    attendance, created = Attendance.objects.get_or_create(user=user, grading_period=grading_period)
    attendance.days_present = days_present
    attendance.save()

    return JsonResponse({'success': True, 'message': 'Days present updated successfully.'})